"""
Convert supporting information to vector, and store in vector database
"""
import math
import os
import re
from typing import (
    List,
    Dict,
    Literal,
    Any,
)
import logging
import cv2
from doc2docx import convert
from docx import Document as read_docx
from docx.document import Document as DocxDocument
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from langchain.docstore.document import Document
from langchain_community.document_loaders.pdf import PDFPlumberLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from markdownify import markdownify
from paddleocr import PPStructure
# from paddleocr import draw_structure_result
from pdf2image import convert_from_path
from weaviate.client import Client
# from weaviate.classes import Filter
from weaviate.util import generate_uuid5
from tqdm import tqdm

# Logger
logger = logging.getLogger("p2v.si2vec")


class Si2Vec:
    """
    Define a pipeline to serialize operation for supporting information in different format, i.e., PDF, DOCX, DOC.
    The operations mainly consist of processing full text and tables into markdown via ocr or docx.oxml and load data
    into Weaviate as batch.
    """

    def __init__(
            self,
            client: Client = None,
    ) -> None:
        """
        Create a new pipeline for supplementary materials.

        Args:
            client: weaviate client
        """
        # TODO: robustness check
        self.client = client

    @staticmethod
    def calculate_distance(bbox1: List[float], bbox2: List[float]) -> float:
        """
        Calculate the distance between the centers of two bounding boxes in pdf page layout.

        Args:
            bbox1: bounding box of the first element
            bbox2: bounding box of the second element
        Returns:
            distance[float]: distance between the centers of two bounding boxes
        """
        center1 = ((bbox1[0] + bbox1[2]) / 2, (bbox1[1] + bbox1[3]) / 2)
        center2 = ((bbox2[0] + bbox2[2]) / 2, (bbox2[1] + bbox2[3]) / 2)
        distance = math.sqrt((center2[0] - center1[0]) ** 2 + (center2[1] - center1[1]) ** 2)
        return distance

    def find_nearest_text(self, table, surrounding_texts):
        """
        Find the nearest text above and below the table.

        Args:
            table: table element
            surrounding_texts: list of text elements
        Returns:
            nearest_text_above: nearest text above the table
            nearest_text_below: nearest text below the table
        """
        table_bbox = table["bbox"]
        min_distance_above = min_distance_below = float('inf')
        nearest_text_above = nearest_text_below = None

        for text in surrounding_texts:
            text_bbox = text["bbox"]
            distance = self.calculate_distance(table_bbox, text_bbox)

            # Check if the text is above or below the table
            if text_bbox[3] < table_bbox[1]:  # Text is above the table
                if distance < min_distance_above:
                    min_distance_above = distance
                    nearest_text_above = "\n ".join([line.get("text") for line in text["res"]])
            elif text_bbox[1] > table_bbox[3]:  # Text is below the table
                if distance < min_distance_below:
                    min_distance_below = distance
                    nearest_text_below = "\n ".join([line.get("text") for line in text["res"]])

        return nearest_text_above, nearest_text_below

    # Search for tables and paragraphs in docx
    def iter_block_items(self, parent):
        """
        Yield each paragraph and table child within *parent*, in document order.
        Each returned value is an instance of either Table or Paragraph. *parent*
        would most commonly be a reference to a main Document object, but
        also works for a _Cell object, which itself can contain paragraphs and tables.
        """

        if isinstance(parent, DocxDocument):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            raise ValueError("something wrong")

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)

            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    # Skip gap between the caption and table
    def table_caption(self, docx, n, direction="backward"):
        """
        Skip empty paragraph, and return caption and footer of the table

        Args:
            docx: docx document
            n: index of paragraph
            direction: forward or backward

        Returns:
            n: index of caption or footer
        """
        if len(docx._body.paragraphs[n].text) == 0 and direction == "backward":
            try:
                return self.table_caption(docx, n - 1, "backward")
            except IndexError:
                return None
        elif len(docx._body.paragraphs[n].text) == 0 and direction == "forward":
            try:
                return self.table_caption(docx, n + 1, "forward")
            except IndexError:
                return None
        else:
            return docx._body.paragraphs[n].text

    def process_docx_table(self, docx: DocxDocument):
        """
        Process tables in docx, and convert them to markdown.
        Note that table and paragraph are elements within the same level in docx.

        Args:
            docx: docx document
        Returns:
            tables[list[dict]]: list of tables, each table is a dict with keys: label, caption, tbody
        """

        tables = []
        num_table = 0
        # Set a cursor for location of elements
        for i, block in enumerate(self.iter_block_items(docx)):
            if isinstance(block, Table):
                table = block
                num_table += 1
                # Caption and footer of table
                caption = self.table_caption(docx, i - num_table, "backward")
                footer = self.table_caption(docx, i - num_table + 1, "forward")
                try:
                    caption = caption + " " + footer if (footer and caption is not None) and (
                            footer != caption) else caption
                except TypeError:
                    caption = "None"
                # Label of table
                label = f"Table S{num_table}"
                # Extract cell content from the table.
                table_md = ""
                for row in table.rows:
                    # Add a "|" at the beginning of each row
                    table_md += "|"
                    for cell in row.cells:
                        if cell.text == "":
                            table_md += " " + "|"
                        else:
                            table_md += cell.text.replace('\n', ' ') + "|"
                    table_md += "\n"
                tbody = table_md
                tables.append({"label": label, "caption": caption, "tbody": tbody})
            else:
                pass
        return tables

    def pdf_handler(self, full_path: str, doi: str):
        """
        Handle pdf file, to extract full text into a series of Document objects and convert tables to markdown.

        Args:
            full_path[str]: full path of the pdf file in local machine

        Returns:
            docs[list[Document]]: a list of full text chunks as Document object
            tables[list[dict]]: list of tables, each table is a dict with keys: label, caption, tbody
        """
        loader = PDFPlumberLoader(full_path)

        # Split full text into smaller chunks
        # Configure TextSplitter
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=512,
                                                       chunk_overlap=100,
                                                       separators=["\n\n", "\n", "(?<=\\. )", "(?<=\\, )"],
                                                       length_function=len)
        # Further split paragraphs in Document objects into smaller chunks using TextSplitter
        docs = text_splitter.split_documents(loader.load())  # List of Documents objects
        # Clean metadata, only keep page number, and replace "\n" with space in the middle of the sentence with re
        docs = [Document(page_content=re.sub(r'(?<!\.)\s*\n\s*(?!\.)', ' ', doc.page_content),
                         metadata={'header_1': 'supplementary material'}) for doc in
                docs]
        # print(docs)

        # # Traditional table extraction with camelot
        # cv = Converter(full_path)
        # output_path = full_path.replace(".pdf", ".docx")
        # cv.convert(output_path)
        #
        # tables = camelot.read_pdf(full_path, flavor='lattice', pages='all', flag_size=True, backend='poppler')
        # for table in tables:
        #     # Convert table to markdown
        #     table_md = table.df.to_markdown()
        #     print(table_md)

        # New method to handle tables in PDF via PaddleOCR

        # Check if output path exists
        output_path = f'../../output/si/{doi.replace("/", "_")}/{os.path.basename(full_path).replace(".pdf", "")}'
        if not os.path.exists(output_path):
            os.makedirs(output_path)

        # Convert pdf to image
        image_paths = convert_from_path(full_path,
                                        output_folder=output_path,
                                        fmt='jpeg',
                                        output_file='page_',
                                        paths_only=True)

        # Replace the model you desire from https://paddlepaddle.github.io/PaddleOCR/main/en/version2.x/ppstructure/models_list.html
        required_paths = {
            "layout_model_dir": r'../../model/paddleocr/layout/ppyolov2_r50vd_dcn_365e_publaynet',
            "table_model_dir": r'../../model/paddleocr/table/en_ppstructure_mobile_v2.0_SLANet_infer',
            "det_model_dir": r'../../model/paddleocr/det/en/en_ppocr_mobile_v2.0_table_det',
            "rec_model_dir": r'../../model/paddleocr/rec/en/en_ppocr_mobile_v2.0_table_rec',
        }

        # Check existence of each path
        missing = []
        for name, path in required_paths.items():
            if not os.path.exists(path):
                missing.append(f"{name}: {path}")

        # Raise error if any path is missing
        if missing:
            msg = "\n".join(missing)
            raise FileNotFoundError(f"The following PaddleOCR model directories were not found:\n{msg}\n\n"
                                    f"Please make sure you have manually downloaded and extracted all required models.")

        table_engine = PPStructure(
            show_log=False,
            image_orientation=True,
            lang="en",
            layout_model_dir=required_paths["layout_model_dir"],
            table_model_dir=required_paths["table_model_dir"],
            det_model_dir=required_paths["det_model_dir"],
            rec_model_dir=required_paths["rec_model_dir"],
        )
        num_table = 0
        tables = []
        # Iterate over pdf images
        for i, img_path in enumerate(image_paths):
            img = cv2.imread(img_path)

            result = table_engine(img)
            save_folder = os.path.basename(img_path).split('.')[0]
            # save_structure_res(result, output_path, save_folder, i)

            # Save the result of structure analysis
            # font_path = '../../model/paddleocr/doc/fonts/latin.ttf'  # Fonts from PaddleOCR
            # image = Image.open(img_path).convert('RGB')
            # im_show = draw_structure_result(image, result, font_path=font_path)
            # im_show = Image.fromarray(im_show)
            # im_show.save(f'{output_path}/{save_folder}/result_page_{i}.jpg')

            # Extract table content into markdown and detect caption
            table_in_page = [element for element in result if element.get('type') == 'table']
            text_in_page = [element for element in result if element.get('type') == 'text']
            for table in table_in_page:
                num_table += 1
                label = f"Table S{num_table}"
                # Get table element
                html = table.get('res').get('html')
                # Convert html to markdown
                tbody = markdownify(html)
                tbl_bbox = table.get('bbox')
                # Find text element closest to the table in the same page
                nearest_text_above, nearest_text_below = self.find_nearest_text(table, text_in_page)
                # Choose text as caption
                caption = nearest_text_above if nearest_text_above is not None else nearest_text_below
                tables.append({"label": label, "caption": caption, "tbody": tbody})

        return docs, tables

    def docx_handler(self, full_path: str):
        """
        Handle docx file, to extract full text into a series of Document objects and convert tables to markdown.
        Args:
            full_path[str]: full path of the docx file in local machine

        Returns:
            docs[list[Document]]: a list of full text chunks as Document object
            tables[list[dict]]: list of tables, each table is a dict with keys: label, caption, tbody
        """
        docx = read_docx(full_path)
        docs = []
        start, end = 0, len(docx.paragraphs)
        for i, paragraph in enumerate(docx.paragraphs):
            docs.append(Document(page_content=paragraph.text.strip(), metadata={'header_1': 'supplementary material'}))
            # TODO: Locate authors and affiliation in the front of paper

            # Locate references
            if "Reference" in paragraph.text.strip() and len(paragraph.text.strip()) < 15:
                end = i

        # Chunk the document based on the ending tag
        docs = docs[start:end]

        # Clear empty string
        docs = [doc for doc in docs if doc.page_content != ""]
        # Tables
        tables = self.process_docx_table(docx)
        return docs, tables

    def doc_handler(self, full_path: str):
        # Convert doc to docx in the same directory
        convert(full_path)
        # Modify file path
        file_name = os.path.basename(full_path)
        file_name = file_name.replace(".doc", ".docx")
        new_path = os.path.join(os.path.dirname(full_path), file_name)

        # Process docx file
        docs, tables = self.docx_handler(new_path)
        return docs, tables

    def load_data(self, docs, tables, md_uuid, batch_size):
        """
        Load data into Weaviate as batch.

        Args:
            docs[list[Document]]: a list of full text chunks as Document object
            tables[list[dict]]: list of tables, each table is a dict with keys: label, caption, tbody
            md_uuid: uuid of metadata object
            batch_size: batch size for loading data into Weaviate
        Returns:
            None
        """
        # Configure batch
        self.client.batch.configure(batch_size=batch_size,
                                    retry_failed_objects=True,
                                    dynamic=True)
        # Populate si data into weaviate associated with their metadata
        with self.client.batch as batch:
            # count_doc = 0
            # count_tbl = 0
            # Add document object to batch
            for doc in tqdm(docs, desc="Loading docs to Weaviate", ncols=100, total=len(docs)):
                properties = {
                    "text": doc.page_content
                }
                # Adding section where the text belongs
                if doc.metadata.keys() is not None:
                    for key in doc.metadata.keys():
                        properties[key] = doc.metadata[key]
                else:
                    pass
                # Return uuid of the object
                rt_uuid = batch.add_object(properties=properties,
                                           collection="FullText",
                                           uuid=generate_uuid5(properties))
                # Cross-reference metadata and SI full text
                batch.add_reference(from_object_collection="FullText",
                                    from_object_uuid=rt_uuid,
                                    from_property_name="hasMetadata",
                                    to_object_uuid=md_uuid,
                                    to_object_collection="Metadata")
                batch.add_reference(from_object_collection="Metadata",
                                    from_object_uuid=md_uuid,
                                    from_property_name="hasFullText",
                                    to_object_uuid=rt_uuid,
                                    to_object_collection="FullText")
                # # Calculate and print progress bar
                # count_doc += 1
                # print(f"Full text loading progress: {count_doc}/{len(docs)}")

            # Populate data into table collection
            for table in tqdm(tables, desc="Loading tables to Weaviate", ncols=100, total=len(tables)):
                # Return uuid of the table in SI
                tbl_uuid = batch.add_object(properties=table,
                                            collection="Table",
                                            uuid=generate_uuid5(table))
                # Cross-reference metadata and table
                batch.add_reference(from_object_collection="Table",
                                    from_object_uuid=tbl_uuid,
                                    from_property_name="hasMetadata",
                                    to_object_uuid=md_uuid,
                                    to_object_collection="Metadata")
                batch.add_reference(from_object_collection="Metadata",
                                    from_object_uuid=md_uuid,
                                    from_property_name="hasTable",
                                    to_object_uuid=tbl_uuid,
                                    to_object_collection="Table")
                # # Calculate and print progress
                # count_tbl += 1
                # print(f"Table loading progress: {count_tbl}/{len(tables)}")

            batch.flush()

    def run(self, doi, paths, root, md_uuid, batch_size=50):
        """
        Run the pipeline.
        """
        num_si = 0
        docs = []
        tables = []
        logger.info("=" * 15 + f"Start processing supplementary material of {doi}..." + "=" * 15)
        for path in paths:
            num_si += 1
            logger.info("-" * 10 + "Supplementary material No.{}...".format(num_si) + "-" * 10)
            full_path = os.path.join(root, path)
            # Determine file type
            try:
                if full_path.endswith(".pdf"):
                    docs, tables = self.pdf_handler(full_path, doi)
                elif full_path.endswith(".docx"):
                    docs, tables = self.docx_handler(full_path)
                elif full_path.endswith(".doc"):
                    docs, tables = self.doc_handler(full_path)
            except TypeError:
                raise "Unsupported file type!"
            except FileNotFoundError:
                logger.info(f"File {full_path} not found!")
            except Exception as e:
                logger.error(f"Error: {e}")
            else:
                # Load data into Weaviate
                self.load_data(docs, tables, md_uuid, batch_size)
                logger.info("-" * 10 + f"Loading supplementary material No.{num_si} data to Weaviate complete!" + "-" * 10)
            finally:
                pass
        logger.info("=" * 20 + f"Finish processing supplementary material of {doi}" + "=" * 20)
